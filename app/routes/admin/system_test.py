"""
System Test Plan Admin Routes
Admin interface for managing system test plans and test execution
"""
from flask import Blueprint, render_template, request, redirect, url_for, flash, jsonify, current_app, abort
from flask_login import login_required, current_user
from datetime import datetime, timezone, timedelta
import json
from app.models.system_test_plan import SystemTestPlan
from app.models.user import User

system_test_bp = Blueprint('system_test', __name__, url_prefix='/admin/system-test')

def admin_required(f):
    """Decorator to require admin access."""
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.is_admin:
            abort(403)  # Forbidden
        return f(*args, **kwargs)
    decorated_function.__name__ = f.__name__
    return decorated_function

@system_test_bp.route('/')
@login_required
@admin_required
def dashboard():
    """System test dashboard with overview statistics."""
    try:
        # Get summary statistics
        test_summary = SystemTestPlan.get_test_summary()
        modules_summary = SystemTestPlan.get_modules_summary()
        
        # Get recent test executions
        recent_tests = SystemTestPlan.get_all(order_by='updated_at desc')[:10]
        
        # Get failed tests requiring attention
        failed_tests = SystemTestPlan.get_all(filters={'test_status': 'failed'})
        
        return render_template('admin/system-test/dashboard.html',
                             test_summary=test_summary,
                             modules_summary=modules_summary,
                             recent_tests=recent_tests,
                             failed_tests=failed_tests)
    except Exception as e:
        current_app.logger.error(f'Error loading dashboard: {str(e)}')
        abort(500)

@system_test_bp.route('/test-plans')
@login_required
@admin_required
def test_plans_list():
    """List all test plans with filtering options."""
    try:
        # Get filter parameters
        module_filter = request.args.get('module', '')
        status_filter = request.args.get('status', '')
        priority_filter = request.args.get('priority', '')
        category_filter = request.args.get('category', '')
        
        # Build filters
        filters = {}
        if module_filter:
            filters['module_name'] = module_filter
        if status_filter:
            filters['test_status'] = status_filter
        if priority_filter:
            filters['priority'] = priority_filter
        if category_filter:
            filters['category'] = category_filter
        
        # Get test plans
        test_plans = SystemTestPlan.get_all(filters=filters)
        
        # Get unique values for filter dropdowns
        all_plans = SystemTestPlan.get_all()
        modules = sorted(list(set(plan.module_name for plan in all_plans if plan.module_name)))
        
        return render_template('admin/system-test/test-plans-list.html',
                             test_plans=test_plans,
                             modules=modules,
                             current_filters={
                                 'module': module_filter,
                                 'status': status_filter,
                                 'priority': priority_filter,
                                 'category': category_filter
                             })
    except Exception as e:
        current_app.logger.error(f'Error loading test plans: {str(e)}')
        abort(500)

@system_test_bp.route('/test-plans/new', methods=['GET', 'POST'])
@login_required
@admin_required
def create_test_plan():
    """Create a new test plan."""
    if request.method == 'POST':
        try:
            test_plan = SystemTestPlan(
                test_plan_no=request.form.get('test_plan_no'),
                module_name=request.form.get('module_name'),
                screen_design_ref=request.form.get('screen_design_ref'),
                description=request.form.get('description'),
                scenario=request.form.get('scenario'),
                expected_results=request.form.get('expected_results'),
                procedure=request.form.get('procedure'),
                test_status=request.form.get('test_status', 'pending'),
                priority=request.form.get('priority', 'medium'),
                category=request.form.get('category', 'functional')
            )
            
            if test_plan.save():
                flash('Test plan created successfully!', 'success')
                return redirect(url_for('system_test.test_plans_list'))
            else:
                flash('Error creating test plan.', 'error')
                abort(500)
        except Exception as e:
            current_app.logger.error(f'Error creating test plan: {str(e)}')
            flash(f'Error creating test plan: {str(e)}', 'error')
            abort(500)
    
    return render_template('admin/system-test/test-plan-form.html', 
                         test_plan=None, 
                         action='Create')

@system_test_bp.route('/test-plans/<int:test_plan_id>')
@login_required
@admin_required
def view_test_plan(test_plan_id):
    """View a specific test plan."""
    try:
        test_plan = SystemTestPlan.get_by_id(test_plan_id)
        if not test_plan:
            abort(404)
        
        return render_template('admin/system-test/test-plan-view.html', 
                             test_plan=test_plan)
    except Exception as e:
        current_app.logger.error(f'Error loading test plan: {str(e)}')
        abort(500)

@system_test_bp.route('/test-plans/<int:test_plan_id>/edit', methods=['GET', 'POST'])
@login_required
@admin_required
def edit_test_plan(test_plan_id):
    """Edit an existing test plan."""
    try:
        test_plan = SystemTestPlan.get_by_id(test_plan_id)
        if not test_plan:
            abort(404)
        
        if request.method == 'POST':
            test_plan.test_plan_no = request.form.get('test_plan_no')
            test_plan.module_name = request.form.get('module_name')
            test_plan.screen_design_ref = request.form.get('screen_design_ref')
            test_plan.description = request.form.get('description')
            test_plan.scenario = request.form.get('scenario')
            test_plan.expected_results = request.form.get('expected_results')
            test_plan.procedure = request.form.get('procedure')
            test_plan.test_status = request.form.get('test_status')
            test_plan.priority = request.form.get('priority')
            test_plan.category = request.form.get('category')
            
            if test_plan.save():
                flash('Test plan updated successfully!', 'success')
                return redirect(url_for('system_test.view_test_plan', test_plan_id=test_plan_id))
            else:
                flash('Error updating test plan.', 'error')
                abort(500)
        
        return render_template('admin/system-test/test-plan-form.html', 
                             test_plan=test_plan, 
                             action='Edit')
    except Exception as e:
        current_app.logger.error(f'Error editing test plan: {str(e)}')
        abort(500)

@system_test_bp.route('/test-plans/<int:test_plan_id>/execute', methods=['GET', 'POST'])
@login_required
@admin_required
def execute_test_plan(test_plan_id):
    """Execute a test plan and record results."""
    try:
        test_plan = SystemTestPlan.get_by_id(test_plan_id)
        if not test_plan:
            abort(404)
        
        if request.method == 'POST':
            test_status = request.form.get('test_status')
            failure_reason = request.form.get('failure_reason', '').strip()
            
            test_plan.test_status = test_status
            test_plan.execution_date = datetime.utcnow()
            test_plan.executed_by = current_user.username
            test_plan.failure_reason = failure_reason if test_status == 'failed' else None
            
            if test_plan.save():
                flash(f'Test execution recorded: {test_status.upper()}', 'success')
                return redirect(url_for('system_test.view_test_plan', test_plan_id=test_plan_id))
            else:
                flash('Error recording test execution.', 'error')
                abort(500)
        
        return render_template('admin/system-test/test-execution.html', 
                             test_plan=test_plan)
    except Exception as e:
        current_app.logger.error(f'Error executing test plan: {str(e)}')
        abort(500)

@system_test_bp.route('/test-plans/<int:test_plan_id>/delete', methods=['POST'])
@login_required
@admin_required
def delete_test_plan(test_plan_id):
    """Delete a test plan."""
    try:
        test_plan = SystemTestPlan.get_by_id(test_plan_id)
        if not test_plan:
            abort(404)
        
        if test_plan.delete():
            flash('Test plan deleted successfully!', 'success')
        else:
            flash('Error deleting test plan.', 'error')
            abort(500)
    except Exception as e:
        current_app.logger.error(f'Error deleting test plan: {str(e)}')
        flash(f'Error deleting test plan: {str(e)}', 'error')
        abort(500)
    
    return redirect(url_for('system_test.test_plans_list'))

@system_test_bp.route('/modules/<module_name>')
@login_required
@admin_required
def module_tests(module_name):
    """View all tests for a specific module."""
    try:
        test_plans = SystemTestPlan.get_by_module(module_name)
        
        # Calculate module statistics
        total = len(test_plans)
        passed = len([t for t in test_plans if t.test_status == 'passed'])
        failed = len([t for t in test_plans if t.test_status == 'failed'])
        pending = len([t for t in test_plans if t.test_status == 'pending'])
        
        module_stats = {
            'total': total,
            'passed': passed,
            'failed': failed,
            'pending': pending,
            'pass_rate': (passed / total * 100) if total > 0 else 0
        }
        
        return render_template('admin/system-test/module-tests.html',
                             module_name=module_name,
                             test_plans=test_plans,
                             module_stats=module_stats)
    except Exception as e:
        current_app.logger.error(f'Error loading module tests: {str(e)}')
        abort(500)

@system_test_bp.route('/bulk-import', methods=['GET', 'POST'])
@login_required
@admin_required
def bulk_import():
    """Bulk import test plans from JSON or CSV."""
    if request.method == 'POST':
        try:
            import_data = request.form.get('import_data')
            if not import_data:
                flash('Please provide import data.', 'error')
                return render_template('admin/system-test/bulk-import.html')
            
            # Parse JSON data
            test_plans_data = json.loads(import_data)
            
            # Validate data structure
            if not isinstance(test_plans_data, list):
                flash('Import data must be a JSON array.', 'error')
                return render_template('admin/system-test/bulk-import.html')
            
            # Import test plans
            imported_count = SystemTestPlan.bulk_import(test_plans_data)
            
            if imported_count > 0:
                flash(f'Successfully imported {imported_count} test plans!', 'success')
                return redirect(url_for('system_test.test_plans_list'))
            else:
                flash('No test plans were imported. Please check your data format.', 'error')
        
        except json.JSONDecodeError:
            flash('Invalid JSON format. Please check your data.', 'error')
        except Exception as e:
            current_app.logger.error(f'Error importing test plans: {str(e)}')
            flash(f'Error importing test plans: {str(e)}', 'error')
    
    return render_template('admin/system-test/bulk-import.html')

@system_test_bp.route('/reports')
@login_required
@admin_required
def reports():
    """Generate test execution reports."""
    try:
        # Get comprehensive test data
        test_summary = SystemTestPlan.get_test_summary()
        modules_summary = SystemTestPlan.get_modules_summary()
        
        # Get failed tests with reasons
        failed_tests = SystemTestPlan.get_all(filters={'test_status': 'failed'})
        
        # Get test execution trends (last 30 days)
        all_tests = SystemTestPlan.get_all()
        
        # Create timezone-aware datetime for comparison
        thirty_days_ago = datetime.now(timezone.utc) - timedelta(days=30)
        
        recent_executions = []
        for test in all_tests:
            if test.execution_date:
                # Ensure execution_date is timezone-aware for comparison
                if test.execution_date.tzinfo is None:
                    # If naive, assume UTC
                    test_date = test.execution_date.replace(tzinfo=timezone.utc)
                else:
                    test_date = test.execution_date
                
                if test_date >= thirty_days_ago:
                    recent_executions.append(test)
        
        return render_template('admin/system-test/reports.html',
                             test_summary=test_summary,
                             modules_summary=modules_summary,
                             failed_tests=failed_tests,
                             recent_executions=recent_executions,
                             report_generated_at=datetime.now())
    except Exception as e:
        current_app.logger.error(f'Error generating reports: {str(e)}')
        abort(500)

# API endpoints for AJAX operations
@system_test_bp.route('/api/test-plans/<int:test_plan_id>/status', methods=['POST'])
@login_required
@admin_required
def update_test_status(test_plan_id):
    """API endpoint to quickly update test status."""
    try:
        test_plan = SystemTestPlan.get_by_id(test_plan_id)
        if not test_plan:
            return jsonify({'success': False, 'error': 'Test plan not found'}), 404
        
        data = request.get_json()
        test_status = data.get('status')
        failure_reason = data.get('failure_reason', '')
        
        if test_status not in ['pending', 'passed', 'failed', 'skipped']:
            return jsonify({'success': False, 'error': 'Invalid status'}), 400
        
        test_plan.test_status = test_status
        test_plan.execution_date = datetime.utcnow()
        test_plan.executed_by = current_user.username
        test_plan.failure_reason = failure_reason if test_status == 'failed' else None
        
        if test_plan.save():
            return jsonify({'success': True})
        else:
            return jsonify({'success': False, 'error': 'Failed to save'}), 500
    
    except Exception as e:
        current_app.logger.error(f'Error updating test status: {str(e)}')
        return jsonify({'success': False, 'error': str(e)}), 500

@system_test_bp.route('/api/modules/summary')
@login_required
@admin_required
def api_modules_summary():
    """API endpoint to get modules summary for charts."""
    try:
        modules_summary = SystemTestPlan.get_modules_summary()
        return jsonify({'success': True, 'data': modules_summary})
    except Exception as e:
        current_app.logger.error(f'Error getting modules summary: {str(e)}')
        return jsonify({'success': False, 'error': str(e)}), 500

@system_test_bp.route('/export/docx')
@login_required
@admin_required
def export_test_plans_docx():
    """Export test plans to DOCX format."""
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from flask import send_file
        import io
        
        # Get all test plans
        test_plans = SystemTestPlan.get_all()
        
        # Create document
        doc = Document()
        
        # Add title
        title = doc.add_heading('System Test Plans Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Set title text color to black
        for run in title.runs:
            run.font.color.rgb = None  # Default black
        
        # Add generation info
        gen_info = doc.add_paragraph(f'Generated on: {datetime.now().strftime("%B %d, %Y at %H:%M")}')
        # Set generation info text color to black
        for run in gen_info.runs:
            run.font.color.rgb = None  # Default black
        doc.add_paragraph('')
        
        # Create one continuous table for all test plans
        if test_plans:
            # Create table with header row + number of test plans
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False  # important to prevent Word from resizing columns

            # Desired column widths
            col_widths = [Inches(1.0), Inches(1.2), Inches(4.0), Inches(0.8)]

            # Apply widths to header row cells
            for i, width in enumerate(col_widths):
                table.columns[i].width = width
                for cell in table.columns[i].cells:
                    cell.width = width

            # Add header row
            header_cells = table.rows[0].cells
            headers = ['Test Plan No.', 'Module', 'Description', 'Remarks']

            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
                # Set text color to black
                header_cells[i].paragraphs[0].runs[0].font.color.rgb = None  # Default black

            # Add data rows for each test plan
            for test_plan in test_plans:
                row_cells = table.add_row().cells

                # Populate row data
                row_data = [
                    test_plan.test_plan_no or 'N/A',
                    test_plan.module_name or 'N/A',
                    test_plan.description or 'N/A',
                    test_plan.failure_reason or 'Passed' if test_plan.test_status == 'passed' else 'N/A'
                ]

                for i, value in enumerate(row_data):
                    row_cells[i].text = str(value)
                    # Force width for each new cell too
                    row_cells[i].width = col_widths[i]
                    # Set text color to black
                    for paragraph in row_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = None  # Default black
        else:
            # Add message if no test plans found
            doc.add_paragraph('No test plans found.')

        # Add spacing at the end
        doc.add_paragraph('')
        
        # Save to BytesIO
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        # Generate filename with current date
        filename = f'system-test-plans-{datetime.now().strftime("%Y-%m-%d")}.docx'
        
        return send_file(
            docx_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        current_app.logger.error(f'Error exporting to DOCX: {str(e)}')
        flash('Error generating DOCX export', 'error')
        return redirect(url_for('system_test.reports'))
